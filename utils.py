import os

import numpy as np
from langdetect import detect
from pypdf import PdfReader
from docx import Document as DocxDocument

from sentence_transformers import SentenceTransformer
from sqlalchemy import select

from models import Dialog, User, Document, DocumentChunk, Embedding


def extract_text(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    if file_path.endswith(".docx"):
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    if file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    raise ValueError("Unsupported file format")


def split_text(text, chunk_size=800):
    chunks = []
    for i in range(0, len(text), chunk_size):
        chunks.append(text[i:i+chunk_size])
    return chunks


model = SentenceTransformer("all-MiniLM-L6-v2")
# model = SentenceTransformer("paraphrase-multilingual-MiniLM-L12-v2")

def create_embedding(text):
    return model.encode(text).tolist()


async def get_user_chat(session, tg_id, username):
    result = await session.execute(select(User).where(User.tg_id == tg_id))
    user = result.scalar_one_or_none()
    if not user:
        user = User(tg_id=tg_id, username=username)
        session.add(user)
        await session.flush()

    # 3️⃣ Получаем диалог
    result = await session.execute(
        select(Dialog)
        .where(Dialog.user_id == user.id)
        .order_by(Dialog.created_at.desc())
        .limit(1)
    )
    dialog = result.scalar_one_or_none()
    if not dialog:
        dialog = Dialog(user_id=user.id)
        session.add(dialog)
        await session.flush()

    return dialog, user


async def process_file(file_path, original_filename, tg_id, username, session):
    file_size = os.path.getsize(file_path)
    file_text = extract_text(file_path)

    dialog = (await get_user_chat(session, tg_id, username))[0]

    doc = Document(
        filename=original_filename,
        file_path=file_path,
        file_size=file_size,
        dialog_id=dialog.id,
        language=detect_lang(file_text[:1000])
    )

    session.add(doc)
    await session.commit()
    await session.refresh(doc)

    # 5️⃣ Извлекаем текст и создаем chunks + embeddings
    chunks = split_text(file_text)

    for i, chunk in enumerate(chunks):
        doc_chunk = DocumentChunk(
            document_id=doc.id,
            text=chunk,
            chunk_index=i,
        )
        session.add(doc_chunk)
        await session.flush()

        vector = create_embedding(chunk)
        embedding = Embedding(
            chunk_id=doc_chunk.id,
            vector=vector,
        )
        session.add(embedding)

    await session.commit()

    return doc


def detect_lang(text: str) -> str:
    try:
        print(text)
        return detect(text[:1000])
    except:
        return "unknown"


def cosine_sim(a, b):
    a, b = np.array(a), np.array(b)
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))